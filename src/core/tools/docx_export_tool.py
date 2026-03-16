import re
import uuid
from pathlib import Path

from pydantic import BaseModel
from docx import Document
from docx.shared import Pt

from src.core.tools.base import ProjectBaseTool


class _DocxExportInput(BaseModel):
    content: str


class DocxExportTool(ProjectBaseTool):
    name: str = "docx_export"
    description: str = (
        "Export markdown text content to a formatted DOCX file. "
        "Supports headings, lists, bold, italic, and inline code. "
        "Returns the file path of the generated DOCX."
    )
    args_schema: type[BaseModel] = _DocxExportInput

    def _parse_inline_markdown(self, paragraph, text: str):
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
            else:
                paragraph.add_run(part)

    async def _arun(self, content: str) -> str:
        output_dir = Path("outputs")
        output_dir.mkdir(exist_ok=True)

        filename = f"form_{uuid.uuid4().hex[:8]}.docx"
        filepath = output_dir / filename

        doc = Document()
        
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        lines = content.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            if stripped.startswith('#'):
                level = len(stripped) - len(stripped.lstrip('#'))
                clean_text = stripped.lstrip('#').strip()
                p = doc.add_heading(level=min(level, 9))
                self._parse_inline_markdown(p, clean_text)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                clean_text = stripped[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                self._parse_inline_markdown(p, clean_text)
            elif stripped.startswith('---'):
                doc.add_paragraph()
            else:
                p = doc.add_paragraph()
                self._parse_inline_markdown(p, stripped)

        doc.save(str(filepath))
        return str(filepath)